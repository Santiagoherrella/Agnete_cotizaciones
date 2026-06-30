"""
resumen_general.py - Ensamblador del reporte final en formato Word.

Propósito general:
Este módulo toma todos los datos extraídos y revisados por los "escuadrones"
(Eléctrico, Mecánico, Accesorios, Logístico) y construye un documento Word (.docx)
ordenado y formateado, listo para ser revisado por los ingenieros de diseño y cotizaciones.

Cuándo usarlo:
Se invoca en `grafo.py` después de que el último escuadrón técnico y logístico termina 
exitosamente (o agota sus intentos). Genera el entregable oficial para la familia procesada.

Requisitos:
- Librería `python-docx` para la generación del documento Word.
- Un estado de grafo (`BotState`) que contenga los datos técnicos `datos_electricos`, 
  `datos_mecanicos`, `datos_accesorios`, `datos_logisticos` y `alertas_diseno`.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from src.schemas.state import BotState
from src.utils.logger import get_logger
logger = get_logger("ResumenGeneral")


def limpiar_texto_xml(texto):
    """
    Elimina caracteres de control invisibles que rompen la librería de Word al guardar el XML subyacente.
    
    Parámetros:
    - texto (str): El texto a procesar.
    
    Retorna:
    - string limpio y compatible con archivos Office Open XML.
    """
    if not isinstance(texto, str):
        return str(texto)
    # Solo permite tabulaciones, saltos de línea y caracteres de texto normales (ASCII/Unicode >= 32)
    return "".join(c for c in texto if ord(c) in (9, 10, 13) or ord(c) >= 32)

def formatear_diccionario(doc: Document, titulo: str, datos: dict):
    """
    Escribe una sección completa en el documento Word extrayendo el valor y detallando el origen.
    Marca con color gris si un dato fue inferido y no leído explícitamente.
    
    Parámetros:
    - doc (Document): Instancia actual del documento python-docx.
    - titulo (str): Encabezado de la nueva sección (e.g. "1. Parámetros Eléctricos").
    - datos (dict): Diccionario tipado (Pydantic volcado) con clave, valor y origen.
    """
    if not datos: return
    
    doc.add_heading(titulo, level=2)
    for clave, info in datos.items():
        if not isinstance(info, dict) or "valor" not in info: continue
        
        nombre_campo = clave.replace("_", " ").title()
        
        # 🛡️ FILTRAMOS EL TEXTO ANTES DE INYECTARLO AL WORD
        valor = limpiar_texto_xml(info.get("valor", "No especificado"))
        origen = limpiar_texto_xml(info.get("origen", "No especificado"))
        
        p = doc.add_paragraph()
        p.add_run(f"• {nombre_campo}: ").bold = True
        p.add_run(f"{valor}")
        
        # Si el dato fue inyectado por una norma o deducido, lo marcamos en gris y cursiva
        if origen not in ["Pliego", "No especificado"]:
            run_origen = p.add_run(f" (Dato asumido por: {origen})")
            run_origen.italic = True
            run_origen.font.color.rgb = RGBColor(128, 128, 128)

def nodo_generar_resumen_tipo(state: BotState):
    """
    Toma los JSON generados en el estado por los 4 escuadrones (Eléctrico, Mecánico, 
    Accesorios y Logístico) y los ensambla en un documento `.docx` para cada familia.
    
    Parámetros:
    - state (BotState): Estado del grafo en LangGraph.
    
    Retorna:
    - Diccionario con la confirmación de la familia procesada y la nueva ruta del Word generado.
    """
    # 1. Identificar la Familia actual
    inventario = state.get("inventario_global", [])
    item_id = state.get("item_actual_id", "")
    trafo_actual = next((t for t in inventario if t["item_id"] == item_id), None)
    
    if not trafo_actual: return {}
    familia = trafo_actual.get("tipo_transformador", "General")
    
    logger.info(f" [Ensamblador Word] Redactando Documento Técnico para la familia: {familia}...")
    
    # 2. Crear el Documento Word
    doc = Document()
    doc.add_heading(f"RESUMEN TÉCNICO Y COMERCIAL - {familia.upper()}", 0)
    
    p_intro = doc.add_paragraph("Documento generado automáticamente a partir de los pliegos del cliente. ")
    p_intro.add_run("Uso exclusivo para Ingeniería de Detalle y Cotización - Magnetron S.A.S.").bold = True
    
    # 3. Inyectar las 4 secciones
    formatear_diccionario(doc, "1. Parámetros Eléctricos", state.get("datos_electricos", {}))
    formatear_diccionario(doc, "2. Datos Mecánicos y Fluidos", state.get("datos_mecanicos", {}))
    formatear_diccionario(doc, "3. Accesorios y Protecciones", state.get("datos_accesorios", {}))
    formatear_diccionario(doc, "4. Logística y Condiciones Comerciales", state.get("datos_logisticos", {}))
    
    # 4. Inyectar las Alertas de Diseño (Banderas Rojas)
    alertas = state.get("alertas_diseno", [])
    
    # Auditoría Completa de Certificaciones para el Word (UL, FM, NTC, etc.)
    try:
        from src.tools.exportador_sdm import auditar_certificaciones, extraer_numero_limpio
        
        logi = state.get("datos_logisticos", {})
        cert_info = logi.get("certificaciones_solicitadas", {})
        cert_val = cert_info.get("valor", "").lower() if isinstance(cert_info, dict) else str(cert_info).lower()
        
        # Extraer variables del trafo actual
        tipo_trafo = trafo_actual.get("tipo_transformador", "")
        fases_str = str(trafo_actual.get("fases", "")).lower()
        fases_num = 1 if any(p in fases_str for p in ["mono", "1", "one", "single", "una"]) else 3
        kva_num = extraer_numero_limpio(trafo_actual.get("potencia", ""))
        
        mec = state.get("datos_mecanicos", {})
        mat_info = mec.get("material_bobinados", "")
        material_bobinado = mat_info.get("valor", "") if isinstance(mat_info, dict) else str(mat_info)
        
        # Ejecutar auditoría
        auditoria = auditar_certificaciones(cert_val, tipo_trafo, fases_num, kva_num, material_bobinado)
        
        # Inyectar Alertas Críticas (Ej. FM)
        alertas.extend(auditoria.get("alertas_criticas", []))
        
        # Inyectar Alerta Normativa si piden UL pero la regla dice que no aplica (Ej. Monofásico)
        if "ul" in cert_val and auditoria.get("ul_aplica") is False:
            alertas.append(f"ALERTA NORMATIVA: El pedido exige certificación UL, pero según las reglas internas NO APLICA para este equipo. Motivo: {auditoria.get('ul_detalle')}.")
            
        # Inyectar Notas (NTC, DOE, Corrugados)
        for nota in auditoria.get("notas_informativas", []):
            alertas.append(f"INFORMACIÓN: {nota}")
            
    except Exception as e:
        logger.info(f" [Ensamblador Word] Error al auditar certificaciones: {e}")
        
    if alertas:
        doc.add_heading("🚨 ALERTAS Y RECOMENDACIONES NORMATIVAS", level=2)
        for alerta in alertas:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ {alerta}")
            run.font.color.rgb = RGBColor(255, 0, 0) # Rojo para alertas
    
    # 5. Guardar Archivo
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    nombre_limpio = "".join([c for c in familia if c.isalpha() or c.isdigit() or c==' ']).rstrip()
    ruta_docx = f"data/outputs/Resumen_{nombre_limpio.replace(' ', '_')}_{timestamp}.docx"
    
    os.makedirs("data/outputs", exist_ok=True)
    doc.save(ruta_docx)
    logger.info(f" [Ensamblador Word] Documento guardado en: {ruta_docx}")
    
    completados_actuales = state.get("resumenes_completados", [])
    rutas_actuales = state.get("rutas_fichas_word", [])
    
    return {
        "resumenes_completados": completados_actuales + [familia],
        "rutas_fichas_word": rutas_actuales + [ruta_docx]
    }

# ==========================================
# METADATA
# tools_used: [python-docx, os, datetime]
# use_cases: [Generación de informes Word, Consolidación de datos JSON a formato lectura]
# reusable_components: [limpiar_texto_xml, formatear_diccionario, nodo_generar_resumen_tipo]
# dependencies: [pip install python-docx]
# ==========================================
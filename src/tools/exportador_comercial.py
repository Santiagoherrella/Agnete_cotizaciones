import os
import pandas as pd
import io
from docx import Document
from docx.shared import Pt
from datetime import datetime
from src.schemas.state import BotState

def nodo_word_comercial(state: BotState):
    print("📝 [Exportador Comercial] Generando Resumen Ejecutivo en Word...")
    
    resumen_texto = state.get("resumen_comercial_ejecutivo", "No hay resumen disponible.")
    tabla_markdown = state.get("tabla_comercial_checklist", "")
    cliente = state.get("cliente_identificado", "Cliente_Generico")
    
    # Crear Documento
    doc = Document()
    
    # Título Principal
    titulo = doc.add_heading('RESUMEN EJECUTIVO COMERCIAL - MAGNETRON USA', 0)
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Cliente: {cliente}")

    # Sección de Resumen Narrativo
    doc.add_heading('1. Análisis del Pliego', level=1)
    doc.add_paragraph(resumen_texto)

    # Guardar
    os.makedirs("data/outputs", exist_ok=True)
    nombre_archivo = f"data/outputs/Resumen_Comercial_{cliente}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    doc.save(nombre_archivo)
    
    print(f"✅ [Word Comercial] Guardado en: {nombre_archivo}")
    return {"rutas_fichas_word": [nombre_archivo]} # Lo añadimos a la lista de entregables

def nodo_excel_comercial(state: BotState):
    print("📊 [Exportador Comercial] Generando Checklist en Excel...")
    
    tabla_md = state.get("tabla_comercial_checklist", "")
    cliente = state.get("cliente_identificado", "Cliente_Generico")
    
    if not tabla_md:
        return {}

    try:
        # Convertimos la tabla Markdown a un DataFrame de Pandas
        # Primero limpiamos el Markdown
        lineas = [l.strip() for l in tabla_md.split('\n') if '|' in l and '---' not in l]
        datos = []
        for l in lineas:
            celdas = [c.strip() for c in l.split('|') if c.strip()]
            if len(celdas) >= 2:
                datos.append(celdas)
        
        if len(datos) > 1:
            df = pd.DataFrame(datos[1:], columns=datos[0])
            
            os.makedirs("data/outputs", exist_ok=True)
            nombre_excel = f"data/outputs/Checklist_Comercial_{cliente}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            
            # Guardar con formato limpio
            df.to_excel(nombre_excel, index=False)
            print(f"✅ [Excel Comercial] Guardado en: {nombre_excel}")
            return {"rutas_tablas_ctg": [nombre_excel]}
            
    except Exception as e:
        print(f"⚠️ Error al generar Excel Comercial: {e}")
        
    return {}